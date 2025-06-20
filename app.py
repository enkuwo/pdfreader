from flask import Flask, request, render_template, send_file, jsonify
import joblib  # For scikit-learn
import fitz  # PyMuPDF
import re
import os
import io
import unicodedata
import json
from collections import Counter
from langdetect import detect
from pdf2image import convert_from_bytes, convert_from_path
from PIL import Image
from fpdf import FPDF
from docx import Document
import pytesseract
from googletrans import Translator
import google.generativeai as genai

# --- Google Gemma 3 API Setup ---
genai.configure(api_key="AIzaSyCIRCrUjUdkRiW0tETGunbe8W1ujleLayQ")
GEMMA_MODEL = "gemma-3-27b-it"

app = Flask(__name__)

STOPWORDS = {
    "the", "and", "to", "of", "in", "a", "is", "it", "that", "on", "for", "with", "as", "was",
    "are", "this", "an", "by", "be", "from", "at", "or", "which", "you", "not", "but", "we"
}

# --- Utilities ---
def extract_text_with_ocr(pdf_bytes):
    images = convert_from_bytes(pdf_bytes.read())
    return [pytesseract.image_to_string(img) for img in images]

def generate_page_images(pdf_path):
    os.makedirs("static/page_images", exist_ok=True)
    for file in os.listdir("static/page_images"):
        os.remove(os.path.join("static/page_images", file))
    images = convert_from_path(pdf_path)
    for i, image in enumerate(images):
        image.save(f'static/page_images/page_{i+1}.png', 'PNG')

def extract_text_from_pdf(pdf):
    doc = fitz.open(stream=pdf.read(), filetype="pdf")
    return [page.get_text() for page in doc]

def detect_language(text_pages):
    try:
        return detect(" ".join(text_pages))
    except:
        return "en"

def gemma_summarize(text, lang="en"):
    prompt = (
        "Summarize the following text clearly and concisely.\n\n"
        f"Text:\n{text[:4000]}\n\nSummary:"
    )
    model = genai.GenerativeModel(GEMMA_MODEL)
    response = model.generate_content(prompt)
    return response.text.strip()

def summarize_full(text_pages, lang):
    return gemma_summarize(" ".join(text_pages), lang)

def summarize_by_page(text_pages, lang):
    return "\n".join(
        f"Page {i+1}:\n{gemma_summarize(page, lang)}\n"
        for i, page in enumerate(text_pages) if len(page.strip()) > 50
    )

def summarize_in_groups(text_pages, lang, group_size=3):
    summaries = []
    for i in range(0, len(text_pages), group_size):
        chunk = " ".join(text_pages[i:i+group_size])
        if len(chunk.strip()) > 50:
            summaries.append(f"Pages {i+1}-{min(i+group_size, len(text_pages))}:\n{gemma_summarize(chunk, lang)}\n")
    return "\n".join(summaries)

def extract_numbers(text_pages):
    return "Found numbers:\n" + ", ".join(re.findall(r'\b\d+(?:\.\d+)?\b', " ".join(text_pages)))

def search_and_summarize_by_keyword(text_pages, keyword, lang):
    keyword = keyword.lower()
    summaries = [f"Page {i+1}:\n{gemma_summarize(page, lang)}\n" for i, page in enumerate(text_pages) if keyword in page.lower()]
    return "\n".join(summaries) if summaries else f"No results found for '{keyword}'."

def find_most_frequent_keyword(text_pages, lang):
    words = re.findall(r'\b[a-zA-Z]{3,}\b', " ".join(text_pages).lower())
    filtered = [word for word in words if word not in STOPWORDS]
    if not filtered:
        return "(No valid keywords found.)"
    top_keyword, _ = Counter(filtered).most_common(1)[0]
    summaries = [f"Page {i+1}:\n{gemma_summarize(page, lang)}\n" for i, page in enumerate(text_pages) if top_keyword in page.lower()]
    return f"Most frequent keyword: **{top_keyword}**\n\n" + "\n".join(summaries)

def ask_question_local(summary, question):
    prompt = (
        "You are a helpful assistant. Based on the summary below, answer the user's question clearly.\n\n"
        f"Summary:\n{summary}\n\n"
        f"Question: {question}\n"
        f"Answer:"
    )
    model = genai.GenerativeModel(GEMMA_MODEL)
    response = model.generate_content(prompt)
    ai_answer = response.text.strip()
    return (
        f"As for your uploaded PDF, here’s what I found:\n\n"
        f"📌 Question: {question}\n"
        f"💬 Answer: {ai_answer}\n\n"
        f"Generally, this is based on the summarized content of the document."
    )

def aichat_main(question):
    prompt = (
        "You are a helpful assistant. Answer the user's question clearly and concisely.\n\n"
        f"Question: {question}\n"
        "Answer:"
    )
    model = genai.GenerativeModel(GEMMA_MODEL)
    response = model.generate_content(prompt)
    ai_answer = response.text.strip()
    return (
        f"Here’s what I found:\n\n"
        f"📌 Question: {question}\n"
        f"💬 Answer: {ai_answer}\n\n"
        f"Generally, this is for educational use only."
    )
def generate_quiz_from_summary(summary, num_questions=5):
    prompt = (
    f"Based on the following summary, generate {num_questions} multiple-choice quiz questions. "
    "Each question should have 4 options labeled A), B), C), D), and provide the correct answer as 'Answer: <Letter>'.\n"
    "Format:\nQ: <question>\nA) <option1>\nB) <option2>\nC) <option3>\nD) <option4>\nAnswer: <Letter>\n\n"
    f"Summary:\n{summary}\n"
)

    model = genai.GenerativeModel(GEMMA_MODEL)
    response = model.generate_content(prompt)
    # Optionally, you can parse the response into a list of (Q, A) pairs for nicer display.
    return response.text.strip()
def parse_quiz_string(quiz_string):
    # Split into questions by double newlines
    questions_raw = re.split(r'\n\s*\n', quiz_string.strip())
    quiz = []
    for qraw in questions_raw:
        # Extract question
        q_match = re.search(r'Q:\s*(.*)', qraw)
        if not q_match:
            continue
        question = q_match.group(1).strip()
        
        # Extract choices (A, B, C, D)
        choices = []
        for letter in ['A', 'B', 'C', 'D']:
            c_match = re.search(rf'{letter}\)\s*(.*)', qraw)
            if c_match:
                choices.append(c_match.group(1).strip())
        
        # Extract correct answer letter
        a_match = re.search(r'Answer:\s*([A-D])', qraw)
        if not a_match:
            continue
        answer_letter = a_match.group(1)
        
        # Map answer letter to actual answer text
        answer_idx = ord(answer_letter) - ord('A')
        correct_answer = choices[answer_idx] if 0 <= answer_idx < len(choices) else ""
        
        quiz.append({
            "question": question,
            "choices": choices,
            "answer": correct_answer
        })
    return quiz


def generate_number_illustrations(text_pages):
    lines = [line.strip() for line in " ".join(text_pages).split('.') if len(line.strip()) > 10 and re.search(r'\d', line)]
    prompt = (
        "Extract number-related facts from the following text. "
        "Format each fact like this: Label in Year = Value\n\n"
        + "\n".join(lines[:50])
    )
    data = []
    try:
        model = genai.GenerativeModel(GEMMA_MODEL)
        response = model.generate_content(prompt)
        results = response.text.strip().split("\n")
        for line in results:
            if "=" in line and "in" in line:
                data.append(line.strip())
    except Exception as e:
        print("⚠️ AI error:", e)
    return data[:10]

def translate_text(text, src_lang='auto', tgt_lang='en'):
    translator = Translator()
    try:
        result = translator.translate(text, src=src_lang, dest=tgt_lang)
        return result.text
    except Exception as e:
        print(f"Translation error: {e}")
        return text

def remove_non_latin(text):
    return unicodedata.normalize('NFKD', text).encode('latin-1', 'ignore').decode('latin-1')

def export_summary_to_pdf(summary_text):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", size=12)
    for line in remove_non_latin(summary_text).split('\n'):
        pdf.multi_cell(0, 10, line)
    output = io.BytesIO()
    pdf.output(output)
    output.seek(0)
    return output

def export_summary_to_word(summary_text):
    doc = Document()
    doc.add_heading('PDF Summary', 0)
    for line in summary_text.split('\n'):
        doc.add_paragraph(line)
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# --- Flask Routes ---
@app.route("/", methods=["GET", "POST"])
def upload_file():
    if request.method == "POST":
        pdf = request.files["file"]
        mode = request.form.get("mode")
        keyword = request.form.get("keyword", "").strip()
        question = request.form.get("question", "").strip()
        chat_question = request.form.get("chat_question", "").strip()
        translate_enabled = request.form.get("translate")
        src_lang = request.form.get("src_lang", "auto").strip()
        tgt_lang = request.form.get("tgt_lang", "en").strip()

        if chat_question and not pdf:
            chat_answer = aichat_main(chat_question)
            return render_template("result.html", summary=None, question=None, answer=None, pages=None, illustrations=None, chat_answer=chat_answer)

        if not pdf and not chat_question:
            return render_template("upload.html", error="Please upload a PDF or ask a question.")

        pdf.save("uploaded.pdf")
        generate_page_images("uploaded.pdf")

        try:
            pdf.seek(0)
            text_pages = extract_text_from_pdf(pdf)
            if not any(text.strip() for text in text_pages):
                raise ValueError("Empty text, using OCR")
        except:
            pdf.seek(0)
            text_pages = extract_text_with_ocr(pdf)

        lang = detect_language(text_pages)

        if translate_enabled:
            try:
                text_pages = [translate_text(page, src_lang, tgt_lang) for page in text_pages]
            except Exception as e:
                return render_template("upload.html", error=f"Translation failed: {str(e)}")
            translated_filename = "translated_summary.txt"
            translated_filepath = os.path.join("static", translated_filename)
            translated_text = "\n\n".join(text_pages)
            with open(translated_filepath, "w", encoding="utf-8") as f:
                f.write(translated_text)
        else:
            translated_filename = None
            translated_text = None

        if mode == "full":
            summary = summarize_full(text_pages, lang)
        elif mode == "page":
            summary = summarize_by_page(text_pages, lang)
        elif mode == "group":
            summary = summarize_in_groups(text_pages, lang)
        elif mode == "numbers":
            summary = extract_numbers(text_pages)
        elif mode == "keyword":
            if not keyword:
                return render_template("upload.html", error="Please enter a keyword.")
            summary = search_and_summarize_by_keyword(text_pages, keyword, lang)
        elif mode == "most_used":
            summary = find_most_frequent_keyword(text_pages, lang)
        elif mode == "illustration":
            illustrations = generate_number_illustrations(text_pages)
            summary = summarize_full(text_pages, lang)

            # Save to static/data.json
            json_data = []
            for item in illustrations:
                if "=" in item:
                    label, value = item.split("=")
                    json_data.append({"label": label.strip(), "value": value.strip()})

            with open("static/data.json", "w", encoding="utf-8") as f:
                json.dump(json_data, f)
            
            chat_answer = aichat_main(chat_question) if chat_question else None
            return render_template("result.html", summary=summary, illustrations=illustrations, question=None, answer=None, pages=None, chat_answer=chat_answer)
        else:
            summary = "(Invalid mode selected.)"

        if question:
            context = summary
            answer = ask_question_local(context, question)
            matched_pages = [i+1 for i, page in enumerate(text_pages) if any(word.lower() in page.lower() for word in answer.split())]
        else:
            answer = None
            matched_pages = None

        chat_answer = aichat_main(chat_question) if chat_question else None

        return render_template("result.html", summary=summary, question=question, answer=answer, pages=matched_pages, illustrations=None, chat_answer=chat_answer, translated_filename=translated_filename, translated_summary=translated_text)

    return render_template("upload.html")

@app.route("/ask", methods=["POST"])
def ask_about_pdf():
    question = request.form.get("question", "").strip()
    with open("uploaded.pdf", "rb") as f:
        text_pages = extract_text_from_pdf(f)
    context = "\n\n".join(text_pages)
    answer = ask_question_local(context, question)
    matched_pages = [i+1 for i, page in enumerate(text_pages) if any(word.lower() in page.lower() for word in answer.split())]
    image_links = [f"/static/page_images/page_{page}.png" for page in matched_pages]
    return render_template("result.html", question=question, answer=answer, pages=matched_pages, images=image_links)
@app.route("/make_quiz", methods=["POST"])
def make_quiz():
    summary = request.form.get("summary", "")
    quiz_string = generate_quiz_from_summary(summary)
    quiz = parse_quiz_string(quiz_string)
    print("QUIZ DATA:", quiz)  # Debugging
    return render_template("quiz.html", quiz=quiz, summary=summary)

@app.route("/download", methods=["POST"])
def download_summary():
    summary = request.form.get("summary")
    file_type = request.form.get("filetype", "pdf")
    if file_type == "word":
        output = export_summary_to_word(summary)
        return send_file(output, as_attachment=True, download_name="summary.docx", mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    else:
        output = export_summary_to_pdf(summary)
        return send_file(output, as_attachment=True, download_name="summary.pdf", mimetype="application/pdf")

if __name__ == "__main__":
    app.run(debug=False)
